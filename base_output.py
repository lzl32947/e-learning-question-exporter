import collections
import re
from abc import ABC, abstractmethod
from typing import List, Tuple, Dict
import openpyxl
from process.base_process import BaseProcess
from questions import BaseQuestion, SingleChoiceQuestion, MultipleChoiceQuestion, JudgeQuestion, \
    CalculationQuestion
from docx import Document


class BaseOutputExcel(ABC):
    def __init__(self, title: str):
        # Output
        self.question_list: List[BaseQuestion] = []
        self.output_name = title
        self.output_list: List[Tuple[str]] = []

    @abstractmethod
    def to_tuple(self, question: BaseQuestion):
        pass

    def write(self):
        bk = openpyxl.Workbook()
        sh1 = bk.active
        sh1.append(("题干（必填）", "题型 （必填）", "选项 A", "选项 B", "选项 C", "选项 D",
                    "选项E(勿删)", "选项F(勿删)", "选项G(勿删)", "选项H(勿删)", "正确答案（必填）", "解析（勿删）",
                    "章节（勿删）", "难度"))
        for item in self.output_list:
            sh1.append(item)
        bk.save("output/{}.xlsx".format(self.output_name))


class BaseOutputWord(ABC):
    def __init__(self, title: str):
        # Output
        self.question_dict: Dict = collections.defaultdict(list)
        self.output_name = title

    @staticmethod
    def write_image(document, content: str):
        p = re.compile(r"(<<<image\d{4}>>>)")
        cl = p.split(content)
        for i in range(len(cl)):
            if re.match(r"<<<image\d{4}>>>", cl[i]):
                u = cl[i].replace("<<<image", "").replace(">>>", "")
                try:
                    document.add_picture("temp/{}.png".format(int(u)))
                except Exception as e:
                    print(f"Error >> {e}")
            else:
                document.add_paragraph(cl[i])

    def write(self):
        save_file = "output/{}.docx".format(self.output_name)
        doc: Document = Document()

        global_index = 1

        for key in self.question_dict.keys():

            start = self.question_dict[key][0]
            if isinstance(start, SingleChoiceQuestion):
                doc.add_paragraph("单选题")
            elif isinstance(start, MultipleChoiceQuestion):
                doc.add_paragraph("多选题")
            elif isinstance(start, JudgeQuestion):
                doc.add_paragraph("判断题")
            elif isinstance(start, CalculationQuestion):
                doc.add_paragraph("简答题")
            for q in self.question_dict[key]:
                self.write_image(doc, "{}、{}".format(global_index, q.question_body))
                if isinstance(q, SingleChoiceQuestion):
                    self.write_image(doc, "A、 {}".format(q.question_choices[0])) if q.question_choices[0] != "" else ""
                    self.write_image(doc, "B、 {}".format(q.question_choices[1])) if q.question_choices[1] != "" else ""
                    self.write_image(doc, "C、 {}".format(q.question_choices[2])) if q.question_choices[2] != "" else ""
                    self.write_image(doc, "D、 {}".format(q.question_choices[3])) if q.question_choices[3] != "" else ""
                    self.write_image(doc, "E、 {}".format(q.question_choices[4])) if q.question_choices[4] != "" else ""
                    self.write_image(doc, "F、 {}".format(q.question_choices[5])) if q.question_choices[5] != "" else ""
                    self.write_image(doc, "G、 {}".format(q.question_choices[6])) if q.question_choices[6] != "" else ""
                    self.write_image(doc, "H、 {}".format(q.question_choices[7])) if q.question_choices[7] != "" else ""
                    self.write_image(doc, "I、 {}".format(q.question_choices[8])) if q.question_choices[8] != "" else ""
                    self.write_image(doc, "正确答案：{}".format(q.question_answer))
                    self.write_image(doc, "解析：{}-> {}".format(q.question_sequence, q.question_explanation))
                elif isinstance(q, MultipleChoiceQuestion):
                    self.write_image(doc, "A、 {}".format(q.question_choices[0])) if q.question_choices[0] != "" else ""
                    self.write_image(doc, "B、 {}".format(q.question_choices[1])) if q.question_choices[1] != "" else ""
                    self.write_image(doc, "C、 {}".format(q.question_choices[2])) if q.question_choices[2] != "" else ""
                    self.write_image(doc, "D、 {}".format(q.question_choices[3])) if q.question_choices[3] != "" else ""
                    self.write_image(doc, "E、 {}".format(q.question_choices[4])) if q.question_choices[4] != "" else ""
                    self.write_image(doc, "F、 {}".format(q.question_choices[5])) if q.question_choices[5] != "" else ""
                    self.write_image(doc, "G、 {}".format(q.question_choices[6])) if q.question_choices[6] != "" else ""
                    self.write_image(doc, "H、 {}".format(q.question_choices[7])) if q.question_choices[7] != "" else ""
                    self.write_image(doc, "I、 {}".format(q.question_choices[8])) if q.question_choices[8] != "" else ""
                    self.write_image(doc, "正确答案：{}".format(q.question_answer))
                    self.write_image(doc, "解析：{}-> {}".format(q.question_sequence, q.question_explanation))
                elif isinstance(q, JudgeQuestion):
                    self.write_image(doc, "正确答案：{}".format("A" if q.question_answer else "B"))
                    self.write_image(doc, "解析：{}-> {}".format(q.question_sequence, q.question_explanation))
                elif isinstance(q, CalculationQuestion):
                    self.write_image(doc, "正确答案：{}".format(q.question_answer))
                    self.write_image(doc, "解析：{}-> {}".format(q.question_sequence, q.question_explanation))
                self.write_image(doc, "章节：{}".format(q.get_chapter()))
                doc.add_paragraph("\n")
                global_index += 1

        doc.save(save_file)


class OutputWord2KaoShiBao(BaseOutputWord):
    def __init__(self, title: str):
        super().__init__(title)

    def append(self, process: BaseProcess, **kwargs):
        for key in process.question_dict.keys():
            for q in process.question_dict[key]:
                self.question_dict[key].append(q)


class OutputExcel2KaoShiBao(BaseOutputExcel):
    def __init__(self, title: str):
        super().__init__(title)

    def append(self, process: BaseProcess, **kwargs):
        for key in process.question_dict.keys():
            for q in process.question_dict[key]:
                self.question_list.append(q)
        for q in self.question_list:
            _ = self.to_tuple(q)
            self.output_list.append(_)
        self.question_list = []

    def to_tuple(self, question: BaseQuestion):
        if isinstance(question, SingleChoiceQuestion):
            return (question.question_body, "单选题",
                    *question.question_choices[:8], question.question_answer,
                    question.question_explanation, question.get_chapter())
        elif isinstance(question, MultipleChoiceQuestion):
            return (question.question_body, "多选题",
                    *question.question_choices[:8], question.question_answer,
                    question.question_explanation, question.get_chapter())
        elif isinstance(question, JudgeQuestion):
            return (question.question_body, "判断题",
                    "正确", "错误", *([""] * 6), "A" if question.question_answer else "B",
                    question.question_explanation, question.get_chapter())
        elif isinstance(question, CalculationQuestion):
            return (question.question_body, "计算题", *([""] * 8),
                    question.question_answer, question.question_explanation, question.get_chapter())
